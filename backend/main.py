"""UniChat AI - Backend"""

import os
import re
import time
import json
import asyncio
import subprocess
import logging
from datetime import datetime
from typing import Optional
from contextlib import asynccontextmanager

from fastapi import FastAPI, HTTPException, WebSocket, WebSocketDisconnect
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse, JSONResponse
from pydantic import BaseModel
import sqlalchemy
from sqlalchemy import create_engine, text
from sqlalchemy.orm import sessionmaker
import httpx

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# ─── Konfiguracija ───────────────────────────────────────────────────────────

MYSQL_URL = os.getenv("MYSQL_URL", "mysql+pymysql://uni_reader:readonly123@localhost:3306/uni_db")
OLLAMA_URL = os.getenv("OLLAMA_URL", "http://localhost:11434")
OLLAMA_MODEL = os.getenv("OLLAMA_MODEL", "uni-chat-qwen")

# ─── Metrike ─────────────────────────────────────────────────────────────────

metrics = {
    "total_queries": 0,
    "successful_queries": 0,
    "total_response_time": 0.0,
}

# ─── SQL validacija ──────────────────────────────────────────────────────────

FORBIDDEN_PATTERNS = [
    r'\b(INSERT|UPDATE|DELETE|DROP|ALTER|CREATE|TRUNCATE|REPLACE|GRANT|REVOKE)\b',
    r'\b(EXEC|EXECUTE|CALL|INTO\s+OUTFILE|INTO\s+DUMPFILE|LOAD_FILE)\b',
    r'\b(INFORMATION_SCHEMA|MYSQL|PERFORMANCE_SCHEMA|SYS)\b',
    r';\s*\w',  # multiple statements
]

# ─── Šeme po ulogama ────────────────────────────────────────────────────────

VALID_TABLES = {"studenti", "profesori", "predmeti", "ocene", "upisi"}
VALID_COLUMNS = {
    "studenti": {"id", "ime", "prezime", "broj_indeksa", "godina_upisa", "smer", "email"},
    "profesori": {"id", "ime", "prezime", "titula", "email", "kabinet"},
    "predmeti": {"id", "naziv", "sifra", "ects", "semestar", "profesor_id"},
    "ocene": {"id", "student_id", "predmet_id", "ocena", "datum_polaganja", "semestar"},
    "upisi": {"id", "student_id", "predmet_id", "akademska_godina", "status"},
}

ROLE_TABLES = {
    "student": {"studenti", "ocene", "upisi", "predmeti"},
    "profesor": {"profesori", "predmeti", "ocene", "studenti"},
    "admin": {"studenti", "profesori", "predmeti", "ocene", "upisi"},
}

SCHEMA_FULL = """
CREATE TABLE studenti (
  id INT PRIMARY KEY,
  ime VARCHAR(50),
  prezime VARCHAR(50),
  broj_indeksa VARCHAR(20),
  godina_upisa INT,
  smer VARCHAR(100),
  email VARCHAR(100)
);

CREATE TABLE profesori (
  id INT PRIMARY KEY,
  ime VARCHAR(50),
  prezime VARCHAR(50),
  titula VARCHAR(50),
  email VARCHAR(100),
  kabinet VARCHAR(20)
);

CREATE TABLE predmeti (
  id INT PRIMARY KEY,
  naziv VARCHAR(100),
  sifra VARCHAR(20),
  ects INT,
  semestar INT,
  profesor_id INT REFERENCES profesori(id)
);

CREATE TABLE ocene (
  id INT PRIMARY KEY,
  student_id INT REFERENCES studenti(id),
  predmet_id INT REFERENCES predmeti(id),
  ocena INT,
  datum_polaganja DATE,
  semestar VARCHAR(20)
);

CREATE TABLE upisi (
  id INT PRIMARY KEY,
  student_id INT REFERENCES studenti(id),
  predmet_id INT REFERENCES predmeti(id),
  akademska_godina VARCHAR(20),
  status VARCHAR(20)
);"""

SCHEMA_STUDENT = """
CREATE TABLE studenti (
  id INT PRIMARY KEY,
  ime VARCHAR(50),
  prezime VARCHAR(50),
  broj_indeksa VARCHAR(20),
  godina_upisa INT,
  smer VARCHAR(100),
  email VARCHAR(100)
);

CREATE TABLE predmeti (
  id INT PRIMARY KEY,
  naziv VARCHAR(100),
  sifra VARCHAR(20),
  ects INT,
  semestar INT,
  profesor_id INT REFERENCES profesori(id)
);

CREATE TABLE ocene (
  id INT PRIMARY KEY,
  student_id INT REFERENCES studenti(id),
  predmet_id INT REFERENCES predmeti(id),
  ocena INT,
  datum_polaganja DATE,
  semestar VARCHAR(20)
);

CREATE TABLE upisi (
  id INT PRIMARY KEY,
  student_id INT REFERENCES studenti(id),
  predmet_id INT REFERENCES predmeti(id),
  akademska_godina VARCHAR(20),
  status VARCHAR(20)
);"""

SCHEMA_PROFESOR = """
CREATE TABLE profesori (
  id INT PRIMARY KEY,
  ime VARCHAR(50),
  prezime VARCHAR(50),
  titula VARCHAR(50),
  email VARCHAR(100),
  kabinet VARCHAR(20)
);

CREATE TABLE predmeti (
  id INT PRIMARY KEY,
  naziv VARCHAR(100),
  sifra VARCHAR(20),
  ects INT,
  semestar INT,
  profesor_id INT REFERENCES profesori(id)
);

CREATE TABLE studenti (
  id INT PRIMARY KEY,
  ime VARCHAR(50),
  prezime VARCHAR(50),
  broj_indeksa VARCHAR(20),
  godina_upisa INT,
  smer VARCHAR(100),
  email VARCHAR(100)
);

CREATE TABLE ocene (
  id INT PRIMARY KEY,
  student_id INT REFERENCES studenti(id),
  predmet_id INT REFERENCES predmeti(id),
  ocena INT,
  datum_polaganja DATE,
  semestar VARCHAR(20)
);"""

def get_system_prompt(role: str, user_id: int = None) -> str:
    if role == "student":
        schema = SCHEMA_STUDENT
        role_line = "Uloga: STUDENT"
        if user_id:
            identity = f"Prijavljeni student: studenti.id = {user_id}. Filtriraj ocene i upise sa WHERE student_id = {user_id}."
        else:
            identity = ""
    elif role == "profesor":
        schema = SCHEMA_PROFESOR
        role_line = "Uloga: PROFESOR"
        if user_id:
            identity = f"Prijavljeni profesor: profesori.id = {user_id}. Filtriraj predmete sa WHERE profesor_id = {user_id}."
        else:
            identity = ""
    else:
        schema = SCHEMA_FULL
        role_line = "Uloga: ADMIN (potpun pristup)"
        identity = ""

    col_ref_lines = []
    for table in sorted(ROLE_TABLES.get(role, VALID_TABLES)):
        if table in VALID_COLUMNS:
            cols = ', '.join(sorted(VALID_COLUMNS[table]))
            col_ref_lines.append(f"  {table}: {cols}")
    col_reference = '\n'.join(col_ref_lines)

    return f"""Ti si univerzitetski AI asistent. Korisnik pita na srpskom.

{role_line}
{identity}

REŽIM RADA:
1. Ako korisnik traži podatke iz baze → odgovori SAMO SQL upitom (bez teksta, bez ```, bez komentara)
2. Ako korisnik postavlja opšte pitanje, traži savet, razgovara, pozdravlja, ili pita bilo šta što NIJE o podacima iz baze → odgovori tekstom na srpskom
   - Započni odgovor sa "TEKST:" pa onda tvoj odgovor
   - Možeš slobodno razgovarati o bilo kojoj temi: studiranje, programiranje, matematika, nauka, saveti, opšta pitanja itd.
   - Budi prijateljski, informativan i detaljan u odgovorima
3. Ako pitanje nije moguće odgovoriti iz baze ali JE o podacima → odgovori sa "TEKST:" i objasni šta možeš, sa primerima

PRAVILA ZA SQL:
- Koristi ISKLJUČIVO tabele i kolone definisane ispod. NE IZMIŠLJAJ tabele ni kolone.
- Samo SELECT upiti. Nikada INSERT/UPDATE/DELETE/DROP/ALTER/CREATE.
- Ocene su od 5 do 10. Ocena 5 = pao, 6-10 = položio.
- Broj indeksa format: '2023/0001'

BAZA — OVO SU JEDINE TABELE I KOLONE KOJE POSTOJE:
{schema}

TAČAN SPISAK KOLONA PO TABELAMA (koristi SAMO ove — ništa drugo):
{col_reference}

NEPOSTOJEĆE KOLONE — NIKAD ne koristi ove:
  korisnik_id NE POSTOJI → koristi student_id
  user_id NE POSTOJI → koristi student_id
  nastavnik_id NE POSTOJI → koristi profesor_id
  predmet NE POSTOJI → koristi naziv (u tabeli predmeti)
  datum NE POSTOJI → koristi datum_polaganja (u tabeli ocene)
  indeks NE POSTOJI → koristi broj_indeksa (u tabeli studenti)
  godina NE POSTOJI → koristi godina_upisa (u tabeli studenti)

DOZVOLJENE TABELE: {', '.join(ROLE_TABLES.get(role, VALID_TABLES))}

PRIMERI:

Pitanje: Prikaži sve studente
SQL: SELECT id, ime, prezime, broj_indeksa, godina_upisa, smer, email FROM studenti ORDER BY prezime

Pitanje: Ocene studenta 2023/0001
SQL: SELECT s.ime, s.prezime, p.naziv, o.ocena, o.datum_polaganja FROM ocene o JOIN studenti s ON o.student_id = s.id JOIN predmeti p ON o.predmet_id = p.id WHERE s.broj_indeksa = '2023/0001'

Pitanje: Moje ocene (student_id = 1)
SQL: SELECT p.naziv, o.ocena, o.datum_polaganja FROM ocene o JOIN predmeti p ON o.predmet_id = p.id WHERE o.student_id = 1

Pitanje: Moj prosek (student_id = 1)
SQL: SELECT ROUND(AVG(o.ocena), 2) AS prosek FROM ocene o WHERE o.student_id = 1

Pitanje: Moji predmeti (student_id = 1)
SQL: SELECT p.naziv, p.sifra, u.akademska_godina, u.status FROM upisi u JOIN predmeti p ON u.predmet_id = p.id WHERE u.student_id = 1

Pitanje: Prosečna ocena po predmetu
SQL: SELECT p.naziv, ROUND(AVG(o.ocena), 2) AS prosecna_ocena, COUNT(*) AS broj_ocena FROM ocene o JOIN predmeti p ON o.predmet_id = p.id GROUP BY p.id, p.naziv ORDER BY prosecna_ocena DESC

Pitanje: Broj studenata po smeru
SQL: SELECT smer, COUNT(*) AS broj_studenata FROM studenti GROUP BY smer ORDER BY broj_studenata DESC

Pitanje: Ko predaje Baze podataka?
SQL: SELECT pr.titula, pr.ime, pr.prezime, pr.email FROM profesori pr JOIN predmeti p ON p.profesor_id = pr.id WHERE p.naziv = 'Baze podataka'

Pitanje: Top 5 studenata po proseku
SQL: SELECT s.ime, s.prezime, s.broj_indeksa, ROUND(AVG(o.ocena), 2) AS prosek FROM ocene o JOIN studenti s ON o.student_id = s.id GROUP BY s.id, s.ime, s.prezime, s.broj_indeksa ORDER BY prosek DESC LIMIT 5

Pitanje: Predmeti na 3. semestru
SQL: SELECT id, naziv, sifra, ects FROM predmeti WHERE semestar = 3

Pitanje: Svi profesori
SQL: SELECT id, ime, prezime, titula, email, kabinet FROM profesori ORDER BY prezime

Pitanje: Upisi za 2024/2025
SQL: SELECT s.ime, s.prezime, s.broj_indeksa, p.naziv, u.status FROM upisi u JOIN studenti s ON u.student_id = s.id JOIN predmeti p ON u.predmet_id = p.id WHERE u.akademska_godina = '2024/2025'

Pitanje: Koliko predmeta ima svaki profesor?
SQL: SELECT pr.ime, pr.prezime, COUNT(p.id) AS broj_predmeta FROM profesori pr JOIN predmeti p ON p.profesor_id = pr.id GROUP BY pr.id, pr.ime, pr.prezime

Pitanje: Koji studenti su na mom predmetu? (profesor_id = 1)
SQL: SELECT DISTINCT s.ime, s.prezime, s.broj_indeksa, s.smer, p.naziv AS predmet FROM upisi u JOIN studenti s ON u.student_id = s.id JOIN predmeti p ON u.predmet_id = p.id WHERE p.profesor_id = 1 ORDER BY p.naziv, s.prezime

Pitanje: Zdravo
TEKST: Zdravo! Ja sam univerzitetski AI asistent. Mogu da pretražim bazu podataka za vas. Probajte na primer:
- "Prikaži sve studente"
- "Prosečna ocena po predmetu"
- "Top 5 studenata po proseku"

Pitanje: Šta mogu da te pitam?
TEKST: Mogu da odgovorim na pitanja o univerzitetskim podacima, ali i da razgovaramo o bilo čemu! Na primer:
- "Koje su moje ocene?" (pretraga baze)
- "Koji studenti su na mom predmetu?" (pretraga baze)
- "Objasni mi šta je rekurzija" (opšte pitanje)
- "Kako da se pripremim za ispit?" (savet)

Pitanje: Objasni mi šta je normalizacija baze podataka
TEKST: Normalizacija baze podataka je proces organizovanja podataka u tabele tako da se smanji redundancija (ponavljanje) i poboljša integritet podataka. Postoji nekoliko normalnih formi:

1. **Prva normalna forma (1NF)** — svako polje sadrži jednu vrednost, bez ponavljajućih grupa
2. **Druga normalna forma (2NF)** — ispunjava 1NF + svaka ne-ključna kolona zavisi od celokupnog primarnog ključa
3. **Treća normalna forma (3NF)** — ispunjava 2NF + nema tranzitivnih zavisnosti

Na primer, naša univerzitetska baza koristi normalizaciju — studenti, predmeti i ocene su u odvojenim tabelama umesto u jednoj velikoj tabeli.

Pitanje: Kako da naučim Python?
TEKST: Evo nekoliko saveta za učenje Pythona:

1. **Počni sa osnovama** — promenljive, tipovi podataka, petlje, uslovi
2. **Vežbaj svaki dan** — piši male programe i rešavaj zadatke
3. **Koristi besplatne resurse** — Python dokumentacija, tutorijalse na YouTube-u
4. **Radi projekte** — napravi kalkulator, to-do aplikaciju, web scraper
5. **Čitaj tuđi kod** — GitHub je odličan izvor primera

Python je odličan izbor jer je čitljiv i ima ogromnu zajednicu!

ODGOVORI SQL UPITOM ili TEKST: odgovorom:"""


# ─── Validacija SQL-a ────────────────────────────────────────────────────────

def validate_sql(sql: str) -> tuple[bool, str]:
    sql_clean = sql.strip().rstrip(';').strip()

    if not sql_clean:
        return False, "Prazan SQL upit."

    if sql_clean.upper().startswith("GREŠKA") or sql_clean.upper().startswith("GRESKA"):
        return False, sql_clean

    if not re.match(r'^\s*(SELECT|WITH)\b', sql_clean, re.IGNORECASE):
        return False, "Dozvoljeni su samo SELECT upiti."

    for pattern in FORBIDDEN_PATTERNS:
        if re.search(pattern, sql_clean, re.IGNORECASE):
            return False, "Upit sadrži nedozvoljene operacije. Dozvoljeni su samo SELECT upiti za čitanje podataka."

    return True, "OK"


def validate_schema(sql: str, role: str) -> tuple[bool, str]:
    sql_upper = sql.upper()

    table_pattern = r'(?:FROM|JOIN)\s+([a-zA-Z_][a-zA-Z0-9_]*)'
    found_tables = set(m.lower() for m in re.findall(table_pattern, sql, re.IGNORECASE))

    sql_keywords_lower = {'select', 'where', 'on', 'and', 'or', 'not', 'as', 'in', 'is',
                          'null', 'true', 'false', 'case', 'when', 'then', 'else', 'end'}
    found_tables -= sql_keywords_lower

    allowed_tables = ROLE_TABLES.get(role, VALID_TABLES)

    invalid_tables = found_tables - VALID_TABLES
    if invalid_tables:
        return False, f"Tabela '{', '.join(invalid_tables)}' ne postoji u bazi. Dostupne tabele: {', '.join(sorted(allowed_tables))}"

    unauthorized_tables = found_tables - allowed_tables
    if unauthorized_tables:
        return False, f"Nemate pristup tabeli '{', '.join(unauthorized_tables)}' sa ulogom '{role}'."

    col_pattern = r'([a-zA-Z_][a-zA-Z0-9_]*)\.([a-zA-Z_][a-zA-Z0-9_]*)'
    prefixed_columns = set()  # track columns already validated via table.column
    for table_ref, col_ref in re.findall(col_pattern, sql):
        table_ref_lower = table_ref.lower()
        col_ref_lower = col_ref.lower()
        prefixed_columns.add(col_ref_lower)

        alias_pattern = rf'(?:FROM|JOIN)\s+([a-zA-Z_][a-zA-Z0-9_]*)\s+(?:AS\s+)?{re.escape(table_ref)}\b'
        alias_match = re.search(alias_pattern, sql, re.IGNORECASE)

        if alias_match:
            actual_table = alias_match.group(1).lower()
        elif table_ref_lower in VALID_COLUMNS:
            actual_table = table_ref_lower
        else:
            continue

        if actual_table in VALID_COLUMNS and col_ref_lower not in VALID_COLUMNS[actual_table]:
            valid_cols = ', '.join(sorted(VALID_COLUMNS[actual_table]))
            return False, f"Kolona '{col_ref}' ne postoji u tabeli '{actual_table}'. Dostupne kolone: {valid_cols}"

    all_valid_columns = set()
    for table in allowed_tables:
        if table in VALID_COLUMNS:
            all_valid_columns.update(VALID_COLUMNS[table])

    SQL_KEYWORDS = {
        'select', 'from', 'where', 'join', 'left', 'right', 'inner', 'outer', 'full',
        'on', 'and', 'or', 'not', 'as', 'in', 'is', 'null', 'true', 'false',
        'group', 'order', 'by', 'having', 'limit', 'offset', 'distinct', 'asc', 'desc',
        'between', 'like', 'exists', 'case', 'when', 'then', 'else', 'end',
        'union', 'all', 'with', 'cross', 'natural', 'using',
        'count', 'sum', 'avg', 'min', 'max', 'round', 'concat',
        'ifnull', 'coalesce', 'cast', 'date', 'year', 'month', 'day',
    }

    known_names = {t.lower() for t in VALID_TABLES}
    for m in re.finditer(r'(?:FROM|JOIN)\s+([a-zA-Z_]\w*)(?:\s+(?:AS\s+)?([a-zA-Z_]\w*))?', sql, re.IGNORECASE):
        known_names.add(m.group(1).lower())
        if m.group(2) and m.group(2).lower() not in SQL_KEYWORDS:
            known_names.add(m.group(2).lower())
    for alias in re.findall(r'\bAS\s+([a-zA-Z_]\w*)', sql, re.IGNORECASE):
        known_names.add(alias.lower())

    all_idents = {w.lower() for w in re.findall(r'\b([a-zA-Z_][a-zA-Z0-9_]*)\b', sql)}
    known_good = SQL_KEYWORDS | known_names | all_valid_columns | prefixed_columns
    unknown = all_idents - known_good

    bad_columns = sorted(ident for ident in unknown if '_' in ident)
    if bad_columns:
        return False, f"Kolona '{bad_columns[0]}' ne postoji u bazi. Dostupne kolone: {', '.join(sorted(all_valid_columns))}"

    return True, "OK"


def clean_llm_sql(raw: str) -> str:
    if not raw:
        return ""

    original = raw.strip()
    logger.info(f"LLM raw output ({len(original)} chars): {original[:200]}...")

    sql = original

    # Strip Qwen 3 <think>...</think> blocks
    sql = re.sub(r'<think>.*?</think>', '', sql, flags=re.DOTALL).strip()

    if '```' in sql:
        import re as _re
        code_blocks = _re.findall(r'```(?:sql|mysql)?\s*\n?(.*?)```', sql, _re.DOTALL | _re.IGNORECASE)
        if code_blocks:
            for block in code_blocks:
                block = block.strip()
                if block.upper().startswith(('SELECT', 'WITH')):
                    sql = block
                    break
            else:
                sql = code_blocks[0].strip()

    sql = sql.strip().strip('`').strip()

    if sql.lower().startswith("sql\n") or sql.lower().startswith("sql "):
        sql = sql[3:].strip()

    upper = sql.upper()
    select_idx = upper.find('SELECT')
    with_idx = upper.find('WITH')

    if select_idx == -1 and with_idx == -1:
        logger.warning(f"No SELECT/WITH found in LLM output")
        return sql

    start_idx = -1
    if select_idx >= 0 and with_idx >= 0:
        start_idx = min(select_idx, with_idx)
    elif select_idx >= 0:
        start_idx = select_idx
    else:
        start_idx = with_idx

    if start_idx > 0:
        sql = sql[start_idx:]

    sql = sql.rstrip(';').strip()

    lines = sql.split('\n')
    sql_lines = []
    sql_starters = {
        'SELECT', 'FROM', 'WHERE', 'JOIN', 'LEFT', 'RIGHT', 'INNER', 'OUTER', 'FULL',
        'ON', 'AND', 'OR', 'GROUP', 'ORDER', 'LIMIT', 'HAVING', 'WITH', 'UNION',
        'CASE', 'WHEN', 'THEN', 'ELSE', 'END', 'NOT', 'IN', 'BETWEEN', 'LIKE',
        'AS', 'DISTINCT', 'CROSS', 'NATURAL', 'USING', 'EXISTS', 'OVER', 'PARTITION',
        'SET', 'OFFSET', 'CONCAT', 'ROUND', 'COUNT', 'SUM', 'AVG', 'MIN', 'MAX',
        'IFNULL', 'COALESCE', 'CAST',
    }
    for i, line in enumerate(lines):
        stripped = line.strip()

        if not stripped:
            if sql_lines:
                break
            continue

        if i == 0:
            sql_lines.append(line)
            continue

        first_word = stripped.split()[0].upper().rstrip('(,')
        is_sql_line = (
            first_word in sql_starters or
            stripped[0] in ('(', ')', ',', "'", '"') or
            '.' in stripped.split()[0] or
            (stripped[0].isdigit() and any(c in stripped for c in ['=', '<', '>', ','])) or
            any(stripped.startswith(c) for c in ['(', ')', ','])
        )

        if is_sql_line:
            sql_lines.append(line)
        else:
            break

    result = '\n'.join(sql_lines).strip().rstrip(';').strip()

    COLUMN_FIXES = {
        r'\bkorisnik_id\b': 'student_id',
        r'\buser_id\b': 'student_id',
        r'\bnastavnik_id\b': 'profesor_id',
    }
    for pattern, replacement in COLUMN_FIXES.items():
        fixed = re.sub(pattern, replacement, result, flags=re.IGNORECASE)
        if fixed != result:
            logger.info(f"Auto-fix: {pattern} → {replacement}")
            result = fixed

    logger.info(f"Cleaned SQL: {result[:200]}")
    return result


# ─── GPU Monitoring ──────────────────────────────────────────────────────────

def get_gpu_stats() -> dict:
    try:
        result = subprocess.run(
            ['nvidia-smi', '--query-gpu=utilization.gpu,memory.used,memory.total,temperature.gpu,power.draw,power.limit,name',
             '--format=csv,noheader,nounits'],
            capture_output=True, text=True, timeout=5
        )
        if result.returncode == 0:
            parts = [p.strip() for p in result.stdout.strip().split(',')]
            if len(parts) >= 7:
                return {
                    "available": True,
                    "gpu_name": parts[6],
                    "utilization": float(parts[0]),
                    "vram_used": float(parts[1]),
                    "vram_total": float(parts[2]),
                    "temperature": float(parts[3]),
                    "power_draw": float(parts[4]),
                    "power_limit": float(parts[5]),
                }
    except (FileNotFoundError, subprocess.TimeoutExpired, Exception) as e:
        logger.warning(f"nvidia-smi nedostupan: {e}")

    import random
    return {
        "available": False,
        "gpu_name": "NVIDIA RTX 5080 (Demo)",
        "utilization": round(random.uniform(15, 85), 1),
        "vram_used": round(random.uniform(4000, 12000), 0),
        "vram_total": 16384.0,
        "temperature": round(random.uniform(45, 75), 0),
        "power_draw": round(random.uniform(80, 320), 1),
        "power_limit": 400.0,
    }


# ─── Database ────────────────────────────────────────────────────────────────

engine = None
SessionLocal = None

def init_db():
    global engine, SessionLocal
    try:
        engine = create_engine(MYSQL_URL, pool_pre_ping=True, pool_size=5, echo=False)
        SessionLocal = sessionmaker(bind=engine)
        logger.info("Konekcija sa bazom uspostavljena.")
    except Exception as e:
        logger.error(f"Greška pri konekciji sa bazom: {e}")
        engine = None
        SessionLocal = None


def execute_query(sql: str) -> tuple[list[str], list[list]]:
    if engine is None:
        raise ConnectionError("MySQL baza nije dostupna. Pokrenite MySQL i seedujte bazu sa: sudo mysql < backend/seed.sql")

    with engine.connect() as conn:
        conn.execute(text("SET SESSION TRANSACTION READ ONLY"))
        result = conn.execute(text(sql))
        columns = list(result.keys())
        rows = [list(row) for row in result.fetchall()]
        return columns, rows


# ─── Ollama LLM ─────────────────────────────────────────────────────────────

async def query_ollama(prompt: str, system_prompt: str, max_tokens: int = 512, is_chat: bool = False) -> str:
    stop_sequences = ["```"]
    if not is_chat:
        stop_sequences += ["Objašnjenje:", "Napomena:", "Ovaj upit", "Rezultat:"]

    try:
        async with httpx.AsyncClient(timeout=120.0) as client:
            response = await client.post(
                f"{OLLAMA_URL}/api/chat",
                json={
                    "model": OLLAMA_MODEL,
                    "messages": [
                        {"role": "system", "content": system_prompt},
                        {"role": "user", "content": prompt},
                    ],
                    "stream": False,
                    "options": {
                        "temperature": 0.3 if is_chat else 0.01,
                        "num_predict": max_tokens,
                        "num_ctx": 4096,
                        "top_p": 0.9,
                        "repeat_penalty": 1.1,
                        "stop": stop_sequences,
                    },
                    "think": False,
                }
            )
            if response.status_code == 200:
                data = response.json()
                msg = data.get("message", {})
                return msg.get("content", "").strip()
            else:
                logger.error(f"Ollama greška: {response.status_code} - {response.text}")
                return None
    except httpx.ConnectError:
        logger.warning("Ollama server nije dostupan, koristim fallback.")
        return None
    except Exception as e:
        logger.error(f"Greška pri komunikaciji sa Ollama: {e}")
        return None



# ─── Export funkcije ─────────────────────────────────────────────────────────

EXPORT_DIR = "/tmp/exports"
os.makedirs(EXPORT_DIR, exist_ok=True)

EXPORT_AVAILABLE = {}

def check_export_availability():
    global EXPORT_AVAILABLE
    try:
        import openpyxl
        EXPORT_AVAILABLE["xlsx"] = True
    except ImportError:
        EXPORT_AVAILABLE["xlsx"] = False
        logger.warning("openpyxl nije instaliran — XLSX export neće biti dostupan")

    try:
        import reportlab
        EXPORT_AVAILABLE["pdf"] = True
    except ImportError:
        EXPORT_AVAILABLE["pdf"] = False
        logger.warning("reportlab nije instaliran — PDF export neće biti dostupan")

    try:
        import docx
        EXPORT_AVAILABLE["docx"] = True
    except ImportError:
        EXPORT_AVAILABLE["docx"] = False
        logger.warning("python-docx nije instaliran — DOCX export neće biti dostupan")

    logger.info(f"Dostupni export formati: {[k for k, v in EXPORT_AVAILABLE.items() if v]}")


def export_xlsx(columns: list[str], rows: list[list], filename: str) -> str:
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side

    wb = Workbook()
    ws = wb.active
    ws.title = "Rezultati upita"

    ws.merge_cells(start_row=1, start_column=1, end_row=1, end_column=max(len(columns), 1))
    title_cell = ws.cell(row=1, column=1, value="Univerzitetski informacioni sistem — Rezultati upita")
    title_cell.font = Font(name='Calibri', size=14, bold=True, color="1F4E79")
    title_cell.alignment = Alignment(horizontal='center')

    ws.cell(row=2, column=1, value=f"Datum izvoza: {datetime.now().strftime('%d.%m.%Y. %H:%M')}")
    ws.cell(row=2, column=1).font = Font(name='Calibri', size=9, italic=True, color="666666")

    header_fill = PatternFill(start_color="1F4E79", end_color="1F4E79", fill_type="solid")
    header_font = Font(name='Calibri', size=11, bold=True, color="FFFFFF")
    thin_border = Border(
        left=Side(style='thin'), right=Side(style='thin'),
        top=Side(style='thin'), bottom=Side(style='thin')
    )

    for col_idx, col_name in enumerate(columns, 1):
        cell = ws.cell(row=4, column=col_idx, value=col_name)
        cell.font = header_font
        cell.fill = header_fill
        cell.alignment = Alignment(horizontal='center')
        cell.border = thin_border

    alt_fill = PatternFill(start_color="F2F7FB", end_color="F2F7FB", fill_type="solid")
    for row_idx, row in enumerate(rows, 5):
        for col_idx, value in enumerate(row, 1):
            cell = ws.cell(row=row_idx, column=col_idx, value=value)
            cell.font = Font(name='Calibri', size=10)
            cell.border = thin_border
            if (row_idx - 5) % 2 == 1:
                cell.fill = alt_fill

    for col_idx in range(1, len(columns) + 1):
        max_length = max(
            len(str(ws.cell(row=4, column=col_idx).value or "")),
            *[len(str(ws.cell(row=r, column=col_idx).value or "")) for r in range(5, 5 + len(rows))]
        ) if rows else len(str(ws.cell(row=4, column=col_idx).value or ""))
        ws.column_dimensions[ws.cell(row=4, column=col_idx).column_letter].width = min(max_length + 4, 40)

    footer_row = 5 + len(rows) + 1
    ws.cell(row=footer_row, column=1, value=f"Укупно записа: {len(rows)}")
    ws.cell(row=footer_row, column=1).font = Font(name='Calibri', size=9, italic=True)

    filepath = os.path.join(EXPORT_DIR, filename)
    wb.save(filepath)
    return filepath


def export_pdf(columns: list[str], rows: list[list], filename: str) -> str:
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A4, landscape
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
    from reportlab.lib.units import cm

    filepath = os.path.join(EXPORT_DIR, filename)
    doc = SimpleDocTemplate(filepath, pagesize=landscape(A4), leftMargin=1.5*cm, rightMargin=1.5*cm)

    styles = getSampleStyleSheet()
    elements = []

    title_style = ParagraphStyle('CustomTitle', parent=styles['Title'], fontSize=16, textColor=colors.HexColor('#1F4E79'))
    elements.append(Paragraph("Univerzitetski informacioni sistem", title_style))
    elements.append(Paragraph(f"Rezultati upita — {datetime.now().strftime('%d.%m.%Y. %H:%M')}", styles['Normal']))
    elements.append(Spacer(1, 0.5*cm))

    table_data = [columns] + [[str(v) for v in row] for row in rows]
    table = Table(table_data, repeatRows=1)

    table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#1F4E79')),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, 0), 9),
        ('FONTSIZE', (0, 1), (-1, -1), 8),
        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
        ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor('#CCCCCC')),
        ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#F2F7FB')]),
        ('TOPPADDING', (0, 0), (-1, -1), 4),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 4),
    ]))

    elements.append(table)
    elements.append(Spacer(1, 0.5*cm))
    elements.append(Paragraph(f"Ukupno zapisa: {len(rows)}", styles['Normal']))

    doc.build(elements)
    return filepath


def export_docx(columns: list[str], rows: list[list], filename: str) -> str:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT

    doc = Document()

    title = doc.add_heading('Univerzitetski informacioni sistem', level=1)
    for run in title.runs:
        run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

    subtitle = doc.add_paragraph(f'Rezultati upita — {datetime.now().strftime("%d.%m.%Y. %H:%M")}')
    subtitle.runs[0].font.size = Pt(10)
    subtitle.runs[0].font.italic = True

    doc.add_paragraph()

    table = doc.add_table(rows=1 + len(rows), cols=len(columns))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, col_name in enumerate(columns):
        cell = table.rows[0].cells[i]
        cell.text = str(col_name)
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Pt(9)

    for row_idx, row in enumerate(rows):
        for col_idx, value in enumerate(row):
            cell = table.rows[row_idx + 1].cells[col_idx]
            cell.text = str(value)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)

    doc.add_paragraph()
    footer = doc.add_paragraph(f'Ukupno zapisa: {len(rows)}')
    footer.runs[0].font.size = Pt(9)
    footer.runs[0].font.italic = True

    filepath = os.path.join(EXPORT_DIR, filename)
    doc.save(filepath)
    return filepath


# ─── FastAPI aplikacija ──────────────────────────────────────────────────────

@asynccontextmanager
async def lifespan(app: FastAPI):
    init_db()
    check_export_availability()
    yield

app = FastAPI(title="Univerzitetski AI Chat Asistent", lifespan=lifespan)

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)


# ─── Pydantic modeli ────────────────────────────────────────────────────────

class ChatRequest(BaseModel):
    message: str
    role: str = "student"
    user_id: Optional[int] = None

class ExportRequest(BaseModel):
    columns: list[str]
    rows: list[list]
    format: str


# ─── Podaci o korisnicima ────────────────────────────────────────────────────

MOCK_STUDENTS = [
    {"id": 1, "label": "Marko Jovanović (2023/0001)", "ime": "Marko", "prezime": "Jovanović", "broj_indeksa": "2023/0001"},
    {"id": 2, "label": "Jelena Petrović (2023/0002)", "ime": "Jelena", "prezime": "Petrović", "broj_indeksa": "2023/0002"},
    {"id": 3, "label": "Nikola Đorđević (2022/0015)", "ime": "Nikola", "prezime": "Đorđević", "broj_indeksa": "2022/0015"},
    {"id": 4, "label": "Ana Nikolić (2023/0003)", "ime": "Ana", "prezime": "Nikolić", "broj_indeksa": "2023/0003"},
    {"id": 5, "label": "Stefan Milić (2022/0008)", "ime": "Stefan", "prezime": "Milić", "broj_indeksa": "2022/0008"},
    {"id": 6, "label": "Milica Popović (2023/0004)", "ime": "Milica", "prezime": "Popović", "broj_indeksa": "2023/0004"},
    {"id": 7, "label": "Luka Stojanović (2022/0012)", "ime": "Luka", "prezime": "Stojanović", "broj_indeksa": "2022/0012"},
    {"id": 8, "label": "Teodora Ilić (2023/0005)", "ime": "Teodora", "prezime": "Ilić", "broj_indeksa": "2023/0005"},
    {"id": 9, "label": "Vuk Marković (2022/0020)", "ime": "Vuk", "prezime": "Marković", "broj_indeksa": "2022/0020"},
    {"id": 10, "label": "Sara Kostić (2023/0006)", "ime": "Sara", "prezime": "Kostić", "broj_indeksa": "2023/0006"},
    {"id": 11, "label": "Đorđe Pavlović (2022/0003)", "ime": "Đorđe", "prezime": "Pavlović", "broj_indeksa": "2022/0003"},
    {"id": 12, "label": "Maja Stanković (2023/0007)", "ime": "Maja", "prezime": "Stanković", "broj_indeksa": "2023/0007"},
    {"id": 13, "label": "Filip Živković (2022/0018)", "ime": "Filip", "prezime": "Živković", "broj_indeksa": "2022/0018"},
    {"id": 14, "label": "Ivana Ristić (2023/0008)", "ime": "Ivana", "prezime": "Ristić", "broj_indeksa": "2023/0008"},
    {"id": 15, "label": "Aleksandar Todorović (2022/0025)", "ime": "Aleksandar", "prezime": "Todorović", "broj_indeksa": "2022/0025"},
    {"id": 16, "label": "Katarina Savić (2023/0009)", "ime": "Katarina", "prezime": "Savić", "broj_indeksa": "2023/0009"},
    {"id": 17, "label": "Nemanja Vasić (2022/0007)", "ime": "Nemanja", "prezime": "Vasić", "broj_indeksa": "2022/0007"},
    {"id": 18, "label": "Mina Radović (2023/0010)", "ime": "Mina", "prezime": "Radović", "broj_indeksa": "2023/0010"},
    {"id": 19, "label": "Petar Đukić (2022/0011)", "ime": "Petar", "prezime": "Đukić", "broj_indeksa": "2022/0011"},
    {"id": 20, "label": "Tamara Lazić (2023/0011)", "ime": "Tamara", "prezime": "Lazić", "broj_indeksa": "2023/0011"},
]

MOCK_PROFESSORS = [
    {"id": 1, "label": "prof. dr Dragan Simić", "ime": "Dragan", "prezime": "Simić", "titula": "Redovni profesor"},
    {"id": 2, "label": "prof. dr Milica Stanković", "ime": "Milica", "prezime": "Stanković", "titula": "Vanredni profesor"},
    {"id": 3, "label": "doc. dr Zoran Pavlović", "ime": "Zoran", "prezime": "Pavlović", "titula": "Docent"},
    {"id": 4, "label": "prof. dr Jelena Mitrović", "ime": "Jelena", "prezime": "Mitrović", "titula": "Redovni profesor"},
]


# ─── API rute ────────────────────────────────────────────────────────────────

@app.get("/api/health")
async def health():
    return {"status": "ok", "message": "Sistem je aktivan"}


@app.get("/api/gpu")
async def gpu_stats():
    return get_gpu_stats()


@app.get("/api/metrics")
async def get_metrics():
    avg_time = (metrics["total_response_time"] / metrics["total_queries"]) if metrics["total_queries"] > 0 else 0
    success_rate = (metrics["successful_queries"] / metrics["total_queries"] * 100) if metrics["total_queries"] > 0 else 0
    return {
        "total_queries": metrics["total_queries"],
        "avg_response_time": round(avg_time, 2),
        "success_rate": round(success_rate, 1),
    }


@app.get("/api/export/formats")
async def get_export_formats():
    formats = []
    if EXPORT_AVAILABLE.get("xlsx"):
        formats.append({"id": "xlsx", "label": "Excel (.xlsx)", "color": "#22c55e", "icon": "table"})
    if EXPORT_AVAILABLE.get("pdf"):
        formats.append({"id": "pdf", "label": "PDF (.pdf)", "color": "#ef4444", "icon": "file"})
    if EXPORT_AVAILABLE.get("docx"):
        formats.append({"id": "docx", "label": "Word (.docx)", "color": "#3b82f6", "icon": "doc"})
    return {"formats": formats}


@app.get("/api/users")
async def get_users():
    return {
        "students": MOCK_STUDENTS,
        "professors": MOCK_PROFESSORS,
    }


# ─── SQL šabloni ────────────────────────────────────────────────────────────

def match_template(message: str, role: str, user_id: int = None) -> str | None:
    msg = message.lower().strip()

    if user_id and role == "student":
        if re.search(r'\bmoj[aei]?\s+ocen[aei]', msg) or re.search(r'\bocen[aei]\b.*\bmoj', msg):
            return f"""SELECT p.naziv, o.ocena, o.datum_polaganja, o.semestar
FROM ocene o JOIN predmeti p ON o.predmet_id = p.id
WHERE o.student_id = {int(user_id)}
ORDER BY o.datum_polaganja DESC"""

        if re.search(r'\bmoj\w*\s+(pros[eč]|prosecn)', msg) or re.search(r'pros[eč].*\bmoj', msg):
            return f"""SELECT ROUND(AVG(o.ocena), 2) AS prosek, COUNT(*) AS broj_ocena
FROM ocene o
WHERE o.student_id = {int(user_id)}"""

        if re.search(r'\bmoj[aei]?\s+(predmet|upis)', msg):
            return f"""SELECT p.naziv, p.sifra, p.ects, u.akademska_godina, u.status
FROM upisi u JOIN predmeti p ON u.predmet_id = p.id
WHERE u.student_id = {int(user_id)}
ORDER BY u.akademska_godina DESC, p.naziv"""

        if re.search(r'polož', msg) and re.search(r'\bmoj|sam\b', msg):
            return f"""SELECT p.naziv, o.ocena, o.datum_polaganja
FROM ocene o JOIN predmeti p ON o.predmet_id = p.id
WHERE o.student_id = {int(user_id)} AND o.ocena >= 6
ORDER BY o.datum_polaganja DESC"""

    if user_id and role == "profesor":
        if re.search(r'\bmoj[aei]?\s+predmet', msg) and not re.search(r'ocen|pros[eč]|student', msg):
            return f"""SELECT p.naziv, p.sifra, p.ects, p.semestar
FROM predmeti p
WHERE p.profesor_id = {int(user_id)}
ORDER BY p.semestar"""

        if re.search(r'pros[eč]|prosecn', msg) and re.search(r'\bmoj|predmet', msg):
            if re.search(r'student|ko\s+ima|najmanj|najvec|najbol|najgor', msg):
                order = "ASC" if re.search(r'najmanj|najgor', msg) else "DESC"
                return f"""SELECT s.ime, s.prezime, s.broj_indeksa, p.naziv AS predmet, ROUND(AVG(o.ocena), 2) AS prosek
FROM ocene o
JOIN studenti s ON o.student_id = s.id
JOIN predmeti p ON o.predmet_id = p.id
WHERE p.profesor_id = {int(user_id)}
GROUP BY s.id, s.ime, s.prezime, s.broj_indeksa, p.id, p.naziv
ORDER BY prosek {order}"""
            return f"""SELECT p.naziv, ROUND(AVG(o.ocena), 2) AS prosecna_ocena, COUNT(*) AS broj_ocena
FROM ocene o
JOIN predmeti p ON o.predmet_id = p.id
WHERE p.profesor_id = {int(user_id)}
GROUP BY p.id, p.naziv
ORDER BY prosecna_ocena DESC"""

        if re.search(r'student', msg) and re.search(r'\bmoj|predmet', msg):
            return f"""SELECT DISTINCT s.ime, s.prezime, s.broj_indeksa, s.smer, p.naziv AS predmet
FROM upisi u
JOIN studenti s ON u.student_id = s.id
JOIN predmeti p ON u.predmet_id = p.id
WHERE p.profesor_id = {int(user_id)}
ORDER BY p.naziv, s.prezime"""

        if re.search(r'ocen', msg) and re.search(r'\bmoj', msg):
            return f"""SELECT s.ime, s.prezime, p.naziv, o.ocena, o.datum_polaganja
FROM ocene o
JOIN studenti s ON o.student_id = s.id
JOIN predmeti p ON o.predmet_id = p.id
WHERE p.profesor_id = {int(user_id)}
ORDER BY p.naziv, s.prezime"""

    has_student = bool(re.search(r'student', msg))
    has_profesor = bool(re.search(r'profesor', msg))
    has_predmet = bool(re.search(r'predmet', msg))
    entity_mentions = sum([has_student, has_profesor, has_predmet])

    if entity_mentions >= 2:
        queries = []
        if has_student:
            queries.append(("Studenti", "SELECT id, ime, prezime, broj_indeksa, godina_upisa, smer, email FROM studenti ORDER BY prezime"))
        if has_profesor:
            queries.append(("Profesori", "SELECT id, ime, prezime, titula, email, kabinet FROM profesori ORDER BY prezime"))
        if has_predmet:
            queries.append(("Predmeti", "SELECT id, naziv, sifra, ects, semestar FROM predmeti ORDER BY semestar, naziv"))
        if queries:
            return ("MULTI", queries)
        return None

    if re.search(r'(svi|sve|prikaži|spisak|lista)\s*(student|studenata|studente)', msg):
        return "SELECT id, ime, prezime, broj_indeksa, godina_upisa, smer, email FROM studenti ORDER BY prezime"

    if re.search(r'(svi|sve|prikaži|spisak|lista)\s*(profesor|profesora|profesore)', msg):
        return "SELECT id, ime, prezime, titula, email, kabinet FROM profesori ORDER BY prezime"

    if re.search(r'(svi|sve|prikaži|spisak|lista)\s*(predmet|predmeta|predmete)', msg):
        return "SELECT id, naziv, sifra, ects, semestar FROM predmeti ORDER BY semestar, naziv"

    if re.search(r'pros[eč].*ocen.*predmet|predmet.*pros[eč].*ocen', msg):
        return """SELECT p.naziv, ROUND(AVG(o.ocena), 2) AS prosecna_ocena, COUNT(*) AS broj_ocena
FROM ocene o JOIN predmeti p ON o.predmet_id = p.id
GROUP BY p.id, p.naziv ORDER BY prosecna_ocena DESC"""

    if re.search(r'top|najbolj', msg) and re.search(r'student|prosek', msg):
        limit = 10
        m = re.search(r'top\s*(\d+)', msg)
        if m:
            limit = min(int(m.group(1)), 50)
        return f"""SELECT s.ime, s.prezime, s.broj_indeksa, ROUND(AVG(o.ocena), 2) AS prosek
FROM ocene o JOIN studenti s ON o.student_id = s.id
GROUP BY s.id, s.ime, s.prezime, s.broj_indeksa
ORDER BY prosek DESC LIMIT {limit}"""

    if re.search(r'(broj|koliko)\s*(student|studenata).*smer|smer.*student', msg):
        return "SELECT smer, COUNT(*) AS broj_studenata FROM studenti GROUP BY smer ORDER BY broj_studenata DESC"

    return None


@app.post("/api/chat")
async def chat(request: ChatRequest):
    start_time = time.time()
    metrics["total_queries"] += 1

    valid_roles = {"student", "profesor", "admin"}
    if request.role not in valid_roles:
        return {
            "success": False,
            "error": f"Nepoznata uloga: '{request.role}'. Dozvoljene uloge: Student, Profesor, Admin.",
            "sql": None, "steps": [], "response_time": 0, "export_formats": [],
        }

    steps = []
    def add_step(name: str, status: str = "active"):
        steps.append({"name": name, "status": status, "timestamp": time.time() - start_time})

    try:
        add_step("Primam upit...", "completed")

        template_sql = match_template(request.message, request.role, request.user_id)
        if template_sql:
            add_step("Prepoznat upit (šablon)...", "completed")

            # Handle multi-query templates (separate tables)
            if isinstance(template_sql, tuple) and template_sql[0] == "MULTI":
                multi_queries = template_sql[1]
                add_step("Izvršavam upite...", "active")
                multi_results = []
                all_sql = []
                for label, q in multi_queries:
                    cols, rows = execute_query(q)
                    serialized = []
                    for row in rows:
                        sr = []
                        for val in row:
                            if hasattr(val, 'isoformat'):
                                sr.append(val.isoformat())
                            elif isinstance(val, bytes):
                                sr.append(val.decode('utf-8', errors='replace'))
                            else:
                                sr.append(val)
                        serialized.append(sr)
                    multi_results.append({"label": label, "columns": cols, "rows": serialized, "row_count": len(serialized)})
                    all_sql.append(f"-- {label}\n{q}")
                steps[-1]["status"] = "completed"
                add_step("Gotovo", "completed")
                elapsed = time.time() - start_time
                metrics["successful_queries"] += 1
                metrics["total_response_time"] += elapsed

                export_formats = []
                if EXPORT_AVAILABLE.get("xlsx"):
                    export_formats.append({"id": "xlsx", "label": "Excel", "color": "#22c55e"})
                if EXPORT_AVAILABLE.get("pdf"):
                    export_formats.append({"id": "pdf", "label": "PDF", "color": "#ef4444"})
                if EXPORT_AVAILABLE.get("docx"):
                    export_formats.append({"id": "docx", "label": "Word", "color": "#3b82f6"})

                return {
                    "success": True,
                    "sql": "\n\n".join(all_sql),
                    "multi_results": multi_results,
                    "columns": None,
                    "rows": None,
                    "row_count": sum(r["row_count"] for r in multi_results),
                    "steps": steps,
                    "response_time": round(elapsed, 2),
                    "export_formats": export_formats,
                }

            sql = template_sql
            logger.info(f"Template match: {sql[:100]}")
        else:
            add_step("Generišem SQL...", "active")

            system_prompt = get_system_prompt(request.role, request.user_id)
            msg_lower = request.message.lower()
            data_keywords = [
                'prikaži', 'prikazi', 'spisak', 'lista', 'svi ', 'sve ',
                'koliko', 'broj', 'top ', 'najbolji',
            ]
            question_keywords = [
                'kako', 'šta je', 'sta je', 'zašto', 'zasto', 'objasni',
                'pomozi', 'savet', 'preporuči', 'preporuci', 'nauči', 'nauci',
                'razlika', 'definicija', 'značenje', 'znacenje', 'misliš', 'mislis',
                'da li', 'možeš li', 'mozes li', 'reci mi', 'kaži mi', 'kazi mi',
                'pripremi', 'pripremim', 'prijavim', 'položim', 'polozim',
            ]
            is_general_chat = any(kw in msg_lower for kw in question_keywords)
            is_data_query = not is_general_chat and any(kw in msg_lower for kw in data_keywords)
            max_tokens = 512 if is_data_query else 1024
            is_chat = not is_data_query
            sql = await query_ollama(request.message, system_prompt, max_tokens=max_tokens, is_chat=is_chat)

            if sql is None:
                steps[-1]["status"] = "error"
                elapsed = time.time() - start_time
                metrics["total_response_time"] += elapsed
                return {
                    "success": False,
                    "error": "Ollama server nije dostupan. Pokrenite Ollama sa: ollama serve",
                    "sql": None,
                    "steps": steps,
                    "response_time": round(elapsed, 2),
                    "export_formats": [],
                }

            raw_response = sql.strip()

            # Detect text responses: explicit TEKST: prefix, or responses that don't look like SQL
            upper_raw = raw_response.upper()
            is_text_response = (
                upper_raw.startswith("TEKST:") or
                upper_raw.startswith("GREŠKA:") or
                upper_raw.startswith("GRESKA:") or
                (is_chat and not upper_raw.startswith(("SELECT", "WITH")))
            )
            if is_text_response:
                if ':' in raw_response and any(raw_response.upper().startswith(p) for p in ("TEKST:", "GREŠKA:", "GRESKA:")):
                    colon_idx = raw_response.index(':')
                    text_response = raw_response[colon_idx + 1:].strip()
                else:
                    text_response = raw_response
                steps[-1]["status"] = "completed"
                add_step("Tekstualni odgovor", "completed")
                elapsed = time.time() - start_time
                metrics["total_response_time"] += elapsed
                metrics["successful_queries"] += 1
                return {
                    "success": True,
                    "text_response": text_response,
                    "sql": None,
                    "columns": None,
                    "rows": None,
                    "row_count": 0,
                    "response_time": round(elapsed, 2),
                    "steps": steps,
                    "export_formats": [],
                }

            sql = clean_llm_sql(sql)
            steps[-1]["status"] = "completed"

            add_step("Validacija upita...", "active")
            is_valid, validation_msg = validate_sql(sql)

            if not is_valid:
                logger.warning(f"SQL validacija neuspešna: {validation_msg}. LLM output: {sql[:200]}. Pokušavam ponovo...")
                retry_prompt = f"""Odgovori ISKLJUČIVO sa jednim SQL SELECT upitom. NIŠTA DRUGO — samo SQL.
Bez teksta pre, bez teksta posle, bez objašnjenja, bez markdown.

Pitanje korisnika: {request.message}

SQL:"""
                retry_sql = await query_ollama(retry_prompt, system_prompt)
                if retry_sql:
                    retry_sql = clean_llm_sql(retry_sql)
                    retry_valid, retry_msg = validate_sql(retry_sql)
                    if retry_valid:
                        sql = retry_sql
                        is_valid = True
                        validation_msg = "OK"
                        logger.info("SQL retry uspešan")

            if not is_valid:
                if validation_msg == "Prazan SQL upit." or (not sql.strip().upper().startswith("SELECT") and not sql.strip().upper().startswith("WITH")):
                    text = raw_response if raw_response and not raw_response.strip().upper().startswith("SELECT") else sql
                    steps[-1]["status"] = "completed"
                    add_step("Tekstualni odgovor", "completed")
                    elapsed = time.time() - start_time
                    metrics["total_response_time"] += elapsed
                    metrics["successful_queries"] += 1
                    return {
                        "success": True,
                        "text_response": text,
                        "sql": None,
                        "columns": None,
                        "rows": None,
                        "row_count": 0,
                        "response_time": round(elapsed, 2),
                        "steps": steps,
                        "export_formats": [],
                    }

                steps[-1]["status"] = "error"
                elapsed = time.time() - start_time
                metrics["total_response_time"] += elapsed
                return {
                    "success": False,
                    "error": validation_msg,
                    "sql": sql,
                    "steps": steps,
                    "response_time": round(elapsed, 2),
                    "export_formats": [],
                }

            schema_valid, schema_msg = validate_schema(sql, request.role)

            if not schema_valid:
                logger.warning(f"Schema greška: {schema_msg}. Pokušavam ponovo...")
                allowed = ROLE_TABLES.get(request.role, VALID_TABLES)
                col_lines = []
                for t in sorted(allowed):
                    if t in VALID_COLUMNS:
                        col_lines.append(f"  {t}: {', '.join(sorted(VALID_COLUMNS[t]))}")
                col_ref = '\n'.join(col_lines)

                retry_prompt = f"""GREŠKA U PRETHODNOM UPITU: {schema_msg}

TAČNE KOLONE:
{col_ref}

VAŽNO: korisnik_id NE POSTOJI. Koristi student_id. user_id NE POSTOJI. Koristi student_id.

Originalno pitanje: {request.message}

Odgovori SAMO SQL upitom. Koristi SAMO kolone iz spiska iznad."""

                retry_sql = await query_ollama(retry_prompt, system_prompt)
                if retry_sql:
                    retry_sql = clean_llm_sql(retry_sql)
                    retry_valid, _ = validate_sql(retry_sql)
                    retry_schema_valid, retry_schema_msg = validate_schema(retry_sql, request.role)

                    if retry_valid and retry_schema_valid:
                        sql = retry_sql
                        schema_valid = True
                        logger.info("Retry uspešan — ispravljen SQL")

                if not schema_valid:
                    steps[-1]["status"] = "error"
                    elapsed = time.time() - start_time
                    metrics["total_response_time"] += elapsed
                    return {
                        "success": False,
                        "error": f"Model je generisao nevažeći SQL. {schema_msg}",
                        "sql": sql,
                        "steps": steps,
                        "response_time": round(elapsed, 2),
                        "export_formats": [],
                    }

            steps[-1]["status"] = "completed"

        add_step("Izvršavam upit...", "active")
        columns, rows = execute_query(sql)
        steps[-1]["status"] = "completed"

        add_step("Formiram odgovor...", "active")

        serialized_rows = []
        for row in rows:
            serialized_row = []
            for val in row:
                if hasattr(val, 'isoformat'):
                    serialized_row.append(val.isoformat())
                elif isinstance(val, bytes):
                    serialized_row.append(val.decode('utf-8', errors='replace'))
                else:
                    serialized_row.append(val)
            serialized_rows.append(serialized_row)

        steps[-1]["status"] = "completed"
        add_step("Gotovo", "completed")

        elapsed = time.time() - start_time
        metrics["successful_queries"] += 1
        metrics["total_response_time"] += elapsed

        export_formats = []
        if EXPORT_AVAILABLE.get("xlsx"):
            export_formats.append({"id": "xlsx", "label": "Excel", "color": "#22c55e"})
        if EXPORT_AVAILABLE.get("pdf"):
            export_formats.append({"id": "pdf", "label": "PDF", "color": "#ef4444"})
        if EXPORT_AVAILABLE.get("docx"):
            export_formats.append({"id": "docx", "label": "Word", "color": "#3b82f6"})

        return {
            "success": True,
            "sql": sql,
            "columns": columns,
            "rows": serialized_rows,
            "row_count": len(serialized_rows),
            "steps": steps,
            "response_time": round(elapsed, 2),
            "export_formats": export_formats,
        }

    except Exception as e:
        elapsed = time.time() - start_time
        metrics["total_response_time"] += elapsed
        logger.error(f"Greška: {e}")
        return {
            "success": False,
            "error": f"Greška pri izvršavanju upita: {str(e)}",
            "sql": sql if 'sql' in dir() else None,
            "steps": steps,
            "response_time": round(elapsed, 2),
            "export_formats": [],
        }


@app.post("/api/export")
async def export_data(request: ExportRequest):
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")

    if not EXPORT_AVAILABLE.get(request.format):
        raise HTTPException(
            status_code=400,
            detail=f"Format '{request.format}' nije dostupan. Potrebna biblioteka nije instalirana."
        )

    try:
        if request.format == "xlsx":
            filepath = export_xlsx(request.columns, request.rows, f"rezultati_{timestamp}.xlsx")
            media_type = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
        elif request.format == "pdf":
            filepath = export_pdf(request.columns, request.rows, f"rezultati_{timestamp}.pdf")
            media_type = "application/pdf"
        elif request.format == "docx":
            filepath = export_docx(request.columns, request.rows, f"rezultati_{timestamp}.docx")
            media_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        else:
            raise HTTPException(status_code=400, detail="Nepoznat format.")

        return FileResponse(
            path=filepath,
            filename=os.path.basename(filepath),
            media_type=media_type,
        )
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Greška pri izvozu: {e}")
        raise HTTPException(status_code=500, detail=f"Greška pri izvozu: {str(e)}")


@app.websocket("/ws/gpu")
async def gpu_websocket(websocket: WebSocket):
    await websocket.accept()
    try:
        while True:
            stats = get_gpu_stats()
            await websocket.send_json(stats)
            await asyncio.sleep(2)
    except WebSocketDisconnect:
        pass


if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000)
